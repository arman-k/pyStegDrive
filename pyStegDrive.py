# pyStegDrive - Google Drive Unlimited Storage
# This program uses steganography to store data in Google Docs, without consuming any quota
# Author: Arman Khandaker

import os, sys, shutil, tempfile # Import for file management
import zlib, base64 # Import for encoding/decoding and compressing/decompressing
from docx import Document # Import python-docx for docx file handling
from pydrive.auth import GoogleAuth # Import pydrive Google Authentication module
from pydrive.drive import GoogleDrive # Import pydrive Google Drive module 

docxType = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' # Mimetype for docx format
folderType = 'application/vnd.google-apps.folder' # Mimetype for Google Drive folder
# Some buffer sizes to be used for processing
compBufSize = 1024
encBufSize = 768
decompBufSize = 1024
maxBufSize = 700000
maxResultSize = 1000

def login(): # Authenticate Google Drive Session
	global gauth, drive
	
	try:
		gauth = GoogleAuth() # Initiate GoogleAuth instance
		if os.path.isfile('credentials'):
			gauth.LoadCredentialsFile('credentials') # Load credentials from file if it exists
			if gauth.access_token_expired == True:
				gauth.Refresh() # If expired, refresh access token
			print '[+]Authentication successful'
		else:
			gauth.LocalWebserverAuth() # Open server for authentication
			gauth.SaveCredentialsFile('credentials') # Save authentication credentials in local file
		drive = GoogleDrive(gauth) # Create GoogleDrive instance
	
	except Exception:
		print '[+]Error authenticating. Make sure you are connected to the internet and login credentials are correct'
		quit()

def upProcess():
	fileName = raw_input('Enter the filename: ')
	print '[+]Processing...'
	compressEncode(fileName) # Compress and encode the file for uploading
	
	try:
		driveFolder = createDriveFolder(fileName)
		driveFolder.Upload() # Upload Google Drive Folder
		
		originalDir = os.getcwd() # Store the current working directory
		folderPath = originalDir + '\\' + fileName + '_pyCF' # Create a _pyCF folder path
		sourceDir = os.listdir(folderPath)
		os.chdir(folderPath)
		
		print '[+]Uploading...'
		upload(sourceDir, driveFolder) # Upload from the source directory to the Drive Folder
		
	except Exception:
		print '[+]Error uploading. Make sure you are connected to the internet and login credentials are correct'
		quit()
	
	os.chdir(originalDir)	
	print '[+]File uploaded successfully'
	
	return folderPath # return path of file to be cleaned up

def compressEncode(src):
	try:
		with open(src, 'rb') as input, tempfile.TemporaryFile() as temp: # Open source file and create temp file for storing the compressed version of the file 
			compress(input, temp) # Compress the file
						
			doc = Document() # Create a docx instance using the python-docx module
			output = doc.add_paragraph()
			temp.seek(0, 0)
			
			dirName = os.getcwd() + '\\' + src + '_pyCF' 
			os.mkdir(dirName) # Make a folder to store the docx chunks
			
			encode(temp, doc, src, dirName, output)
						
	except IOError:
		print '[+]File processing error. Please make sure you have sufficient disk space and try again later'
		quit()

def compress(input, temp):
	comp = zlib.compressobj()
	while True:
		buf = input.read(compBufSize)
		if len(buf) == 0:
			break
		temp.write(comp.compress(buf)) # Compress using zlib module
	
	temp.write(comp.flush()) # Compress any unprocessed input

def encode(temp, doc, src, dirName, output):
	size = 0
	tmpVarId = 1
	
	while True:
		buf = temp.read(encBufSize)
		size += len(buf)
		if size >= maxBufSize:
			size = 0
			docSave(tmpVarId, doc, src, dirName) # Save the chunk
			tmpVarId += 1
			doc = Document() # Create a new doc to store the next chunk
			output = doc.add_paragraph()
		if len(buf) == 0:
			break
		output.add_run(base64.b64encode(buf)) # Encode the content in base64 and append it to the docx file
		output = doc.add_paragraph() # Buffer the content into paragraphs
			
	docSave(tmpVarId, doc, src, dirName) # Save the encoded file
	
def docSave(tmpVarId, doc, src, dirName): 
	encodedFileName = src + str(tmpVarId) + '.docx'
	encodedFilePath = os.path.join(dirName, encodedFileName)
	doc.save(encodedFilePath) # Save the docx file
	
def createDriveFolder(fileName):
	splitted = fileName.split('.')
	justTheName = splitted[0] + '_pyCF_' + splitted[len(splitted)-1]
	
	return drive.CreateFile({'title': justTheName, 'mimeType': folderType}) # Create GoogleDriveFolder instance

def upload(sourceDir, driveFolder):
	for file in sourceDir:
		splitted = file.split('.')
		justTheName = splitted[0] + '.' + splitted[len(splitted)-2]
		driveFile = drive.CreateFile({'title': justTheName, 'parents': [{'id': driveFolder['id']}]}) 
		driveFile.SetContentFile(file)
		driveFile.Upload(param={'convert': True}) # Upload the file and convert it to Google Doc format
		
def downProcess():
	print '[+]Listing files in your Drive...'
	file_list = listFiles()
	folderName = raw_input('Enter the filename you want to download: ') # Choose _pyCF_ folder to download
	
	try:
		for file in file_list:
			if file['title'] == folderName:
				fid = file['id'] # Grab the folder ID for the selected _pyCF_ folder
				
		destDir = os.getcwd() + '\\' + folderName
		os.mkdir(destDir) # Create directory to download the files
		
		print '[+]Downloading...'
		download(fid, destDir) # Download the files
		
		print '[+]Processing...'
		decodeDecompress(folderName) # Decode and reconstruct the original file
	
	except Exception:
		print '[+]Error with connection. Make sure you are connected to the internet'
		quit()
	
	finally:
		cleanup(folderName) # Remove the docx files
		if os.path.isfile(folderName + '_tmp'):
			splitted = folderName.split('_')
			fileName = splitted[0] + '.' + splitted[len(splitted)-1]
			os.rename(folderName + '_tmp', fileName) # Rename the output file from its temporary format
		
	print '[+]File downloaded successfully'
	
def listFiles(): # List all the filenames in the root folder of Google Drive
	try:
		file_list = drive.ListFile({'q': "'root' in parents and trashed=false", 'maxResults': maxResultSize}).GetList() # Get list of files
		for file in file_list:
			print 'Title: %s, ID: %s' % (file['title'], file['id']) # Print the names and IDs
	
	except Exception:
		print '[+]Error connecting. Make sure you are connected to the internet and login credentials are correct'
		quit()
		
	return file_list

def download(fid, destDir):
	for file_list in drive.ListFile({'q': "'%s' in parents and trashed=false"%fid, 'maxResults': maxResultSize}): # Fetch list of files to be downloaded
		for file in file_list:
			driveFile = drive.CreateFile({"parents": [{"id": fid}], 'id': file['id']}) # Create GoogleDriveFile instance for file
			fileName = destDir + '\\' + file['title'] + '.docx'
			driveFile.GetContentFile(fileName, mimetype=docxType) # Download file in docx format
	
def decodeDecompress(srcDir):
	try:
		with open(srcDir + '_tmp', 'wb') as output, tempfile.TemporaryFile() as temp: # Create output file in a temporary format and a temp file to store the data decoded from base64 format
			originalDir = os.getcwd()
			decode(srcDir, temp) # Decode the files in the source directory	
			
			temp.seek(0, 0)
			decompress(temp, output) # Decompress the original file	
			
			os.chdir(originalDir)
		
	except IOError:
		print '[+]File processing error. Please make sure you have sufficient disk space'
		quit()
	
def decode(srcDir, temp):
	os.chdir(srcDir)
	for file in os.listdir(os.getcwd()):
		fileName = os.getcwd() + '\\' + file 
		doc = Document(fileName) # Create a docx instance of the file to be reconstructed using python-docx module
		for p in doc.paragraphs:
			if len(p.text) == 0:
				break
			temp.write(base64.b64decode(p.text)) # Decode the content and append it to temp file	

def decompress(temp, output):
	decomp = zlib.decompressobj()
	while True:
		buf = temp.read(decompBufSize)
		if len(buf) == 0:
			break
		output.write(decomp.decompress(buf)) # Decompress the decoded file using zlib module
		output.write(decomp.decompress(decomp.unconsumed_tail))
	
	output.write(decomp.flush()) # Decompress any unprocessed data
		
def cleanup(temp): # Clean up given folders 
	if os.path.isdir(temp):
		shutil.rmtree(temp)

def main():
	print '[+]Authenticating...'
	login() # Authenticate to your Google Drive account
	while True:
		command = raw_input('Enter u to upload and d to download and any other key to exit: ') # Prompt command
		if command == 'u':
			temp = upProcess()
			cleanup(temp)
		elif command == 'd':
			downProcess()
		else:
			break
	print '[+]Quitting...'
	
if __name__ == '__main__': # Execute main method
	main()